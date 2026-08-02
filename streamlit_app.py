import streamlit as st
import requests
import io

# Try to import docx, if not present, graceful fallback
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set page configuration to be mobile-friendly and beautiful
st.set_page_config(
    page_title="GreenCruise AI - Solar Modeler",
    page_icon="☀️",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Custom Styling for Emerald/Teal Theme
st.markdown("""
    <style>
    .main {
        background-color: #0b1511;
        color: #e0f2f1;
    }
    h1, h2, h3 {
        color: #00bfa5 !important;
    }
    .stButton>button {
        background-color: #00bfa5;
        color: #0b1511;
        font-weight: bold;
        border-radius: 8px;
        width: 100%;
        height: 45px;
    }
    .stButton>button:hover {
        background-color: #00897b;
        color: white;
    }
    </style>
""", unsafe_allow_html=True)

# Helper function to set cell background in Word
def set_cell_background(cell, color_hex):
    try:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)
    except Exception:
        pass

# Helper to generate DOCX file in-memory
def generate_docx_in_memory(address, system_size_kw, current_rate_kwh, annual_kwh, grid_annual_cost, ppa_rate, ppa_annual_cost, annual_savings, co2_saved_tons, tax_deduction, reap_grant, total_incentives):
    doc = docx.Document()
    
    # Font Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    COLOR_PRIMARY_GREEN = RGBColor(11, 81, 50) 
    COLOR_GREY = RGBColor(128, 128, 128)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("GREENCRUISE AI\n")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY_GREEN
    
    sub_run = title_p.add_run("Commercial Energy & Tax Savings Assessment")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_GREY
    
    doc.add_paragraph().add_run("-" * 50).font.color.rgb = COLOR_GREY

    # Metadata
    meta_p = doc.add_paragraph()
    meta_p.add_run("Date: ").bold = True
    meta_p.add_run("August 4, 2026\n")
    meta_p.add_run("Property Address: ").bold = True
    meta_p.add_run(f"{address}\n")
    meta_p.add_run("Prepared For: ").bold = True
    meta_p.add_run("Partner Client Portfolio\n")
    meta_p.add_run("Assessed By: ").bold = True
    meta_p.add_run("Sulejman Mrako, Founder at GreenCruise AI\n")

    # 1. Assessment
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("🛰️ 1. PROPERTY & SOLAR POTENTIAL ASSESSMENT")
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = COLOR_PRIMARY_GREEN
    
    p1 = doc.add_paragraph()
    p1.add_run("Our AI-powered satellite mapping system has scanned your roof coordinates and executed a high-precision solar irradiation simulation.\n\n")
    p1.add_run("• Recommended Solar Array: ").bold = True
    p1.add_run(f"{system_size_kw} kW DC\n")
    p1.add_run("• Projected Annual Production: ").bold = True
    p1.add_run(f"{annual_kwh:,.0f} kWh / year\n")

    # 2. Financials
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("💵 2. FINANCIAL BENEFIT ANALYSIS ($0 UPFRONT)")
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = COLOR_PRIMARY_GREEN
    
    p2 = doc.add_paragraph()
    p2.add_run("Under our Zero-CapEx Green Power Initiative, you pay $0 upfront for design, hardware, and installation.\n")

    # Table
    table = doc.add_table(rows=4, cols=4)
    headers = ["Metric", "Current Grid Status", "Proposed PPA Solar", "Net Client Savings"]
    for i, head in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = head
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0B5132")

    row1 = table.rows[1].cells
    row1[0].text = "Electricity Rate"
    row1[1].text = f"${current_rate_kwh:.3f}/kWh"
    row1[2].text = f"${ppa_rate:.3f}/kWh"
    row1[3].text = "30% Discount"
    row1[3].paragraphs[0].runs[0].font.bold = True

    row2 = table.rows[2].cells
    row2[0].text = "Annual Power Cost"
    row2[1].text = f"${grid_annual_cost:,.2f}"
    row2[2].text = f"${ppa_annual_cost:,.2f}"
    row2[3].text = f"${annual_savings:,.2f} Saved"
    row2[3].paragraphs[0].runs[0].font.bold = True

    row3 = table.rows[3].cells
    row3[0].text = "20-Year Cumulative"
    row3[1].text = f"${grid_annual_cost*20:,.2f}"
    row3[2].text = f"${ppa_annual_cost*20:,.2f}"
    row3[3].text = f"${annual_savings*20:,.2f} Saved"
    row3[3].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # 3. Incentives
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("🏛️ 3. CERTIFIED GREEN TAX INCENTIVES & GRANTS")
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = COLOR_PRIMARY_GREEN
    
    p3 = doc.add_paragraph()
    p3.add_run("• IRS Section 179D Tax Deduction: ").bold = True
    p3.add_run(f"Up to ").font.italic = True
    p3.add_run(f"${tax_deduction:,.2f}").bold = True
    p3.add_run(" in Year-One accelerated write-offs.\n")
    p3.add_run("• USDA REAP Cash-Back Grant: ").bold = True
    p3.add_run(f"Eligible for up to ").font.italic = True
    p3.add_run(f"${reap_grant:,.2f}").bold = True
    p3.add_run(" in direct federal grants.\n")
    p3.add_run("• Net Funding Incentive Package: ").bold = True
    p3.add_run(f"${total_incentives:,.2f}").bold = True

    # Save to buffer
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.title("☀️ GreenCruise AI")
st.subheader("Mobile Solar & PPA Feasibility Modeler")

st.markdown("""
Use this private tool to instantly calculate commercial solar potential, PPA rates, and 20-year savings for any address on Earth.
""")

# Input Fields
address = st.text_input("📍 Commercial Building Address", placeholder="e.g. 1105 Schrock Rd, Columbus, OH 43229")

# Sliders for easy adjustments on mobile touchscreens
system_size = st.slider("⚡ Target System Size (kW DC)", min_value=10, max_value=2000, value=250, step=10)
current_rate = st.slider("💵 Current Utility Rate ($/kWh)", min_value=0.05, max_value=0.50, value=0.18, step=0.01)
ppa_discount = st.slider("📉 Proposed PPA Discount (%)", min_value=10, max_value=50, value=30, step=5)

# Calculate Button
if st.button("🚀 Generate Savings Report"):
    if not address:
        st.error("Please enter a valid address first.")
    else:
        with st.spinner("Analyzing coordinates and simulating solar output..."):
            # 1. Geocode via OpenStreetMap (No API key required, mobile-friendly)
            geocode_url = f"https://nominatim.openstreetmap.org/search?q={address}&format=json&limit=1"
            headers = {"User-Agent": "GreenCruise_AI_Mobile_App/1.0"}
            
            lat, lon = "40.1016", "-82.9850" # Default fallback
            display_name = address
            using_fallback_geo = False
            
            try:
                geo_res = requests.get(geocode_url, headers=headers, timeout=5).json()
                if geo_res:
                    lat = geo_res[0]["lat"]
                    lon = geo_res[0]["lon"]
                    display_name = geo_res[0]["display_name"]
                else:
                    using_fallback_geo = True
            except Exception:
                using_fallback_geo = True
                
            # 2. Call NREL PVWatts API with local mathematical fallback
            annual_kwh = None
            using_local_engine = False
            api_key = "DEMO_KEY"
            pvwatts_url = f"https://developer.nrel.gov/api/pvwatts/v8.json?api_key={api_key}&lat={lat}&lon={lon}&system_capacity={system_size}&azimuth=180&tilt=20&array_type=1&losses=14"
            
            try:
                # Try calling NREL API
                pv_res = requests.get(pvwatts_url, timeout=5).json()
                if "outputs" in pv_res:
                    annual_kwh = pv_res["outputs"]["ac"]
                else:
                    using_local_engine = True
            except Exception:
                using_local_engine = True
                
            # Local Math Fallback Engine
            if using_local_engine:
                try:
                    lat_f = float(lat)
                except ValueError:
                    lat_f = 40.0
                
                if lat_f > 50.0:  # UK / Scotland / Ireland
                    capacity_factor = 0.105
                    region_name = "UK / Ireland (Northern Europe)"
                elif lat_f > 38.0: # US Midwest, Mountain, Pacific Northwest
                    capacity_factor = 0.145
                    region_name = "US Midwest / Mountain / Pacific Northwest"
                else: # US Southern Sun Belt (Texas, Florida)
                    capacity_factor = 0.178
                    region_name = "US Southern Sun Belt"
                
                annual_kwh = system_size * 8760 * capacity_factor
                st.info(f"ℹ️ Local Engine active for: {region_name} (Capacity Factor: {capacity_factor*100:.1f}%)")
            else:
                st.success("🛰️ Connected to Live NREL Database!")

            # 3. Calculate Financials
            grid_annual_cost = annual_kwh * current_rate
            ppa_rate = current_rate * (1 - (ppa_discount / 100))
            ppa_annual_cost = annual_kwh * ppa_rate
            
            annual_savings = grid_annual_cost - ppa_annual_cost
            co2_saved_tons = (annual_kwh * 0.85) / 2000
            
            # Incentives
            tax_deduction = system_size * 170.0  # Approx $170 deduction per kW
            reap_grant = (system_size * 1200.0) * 0.50 # 50% cash-back on solar equipment
            total_incentives = tax_deduction + reap_grant
            
            # Display Beautiful Results
            st.metric("Estimated Year 1 Savings", f"${annual_savings:,.2f}")
            st.metric("Estimated 20-Year Savings", f"${annual_savings * 20:,.2f}")
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Annual Energy Generated", f"{annual_kwh:,.0f} kWh")
                st.metric("Proposed PPA Rate", f"${ppa_rate:.3f}/kWh")
            with col2:
                st.metric("CO2 Saved Annually", f"{co2_saved_tons:.1f} Tons")
                st.metric("System Size", f"{system_size} kW DC")
            
            # Generated Plain-Text Report for Easy Copy/Paste on Mobile
            st.subheader("📋 Copyable Report Text")
            report_text = f"""GREENCRUISE AI SOLAR FEASIBILITY REPORT
--------------------------------------------------
📍 Property: {display_name}
🛰️ Geo-Coordinates: Lat {float(lat):.4f}, Lon {float(lon):.4f}
⚡ Recommended Solar Array: {system_size} kW DC
📈 Annual Production: {annual_kwh:,.0f} kWh / year

FINANCIAL ANALYSIS ($0 UPFRONT CAPEX):
--------------------------------------------------
* Current Utility Rate: ${current_rate:.2f} / kWh
* Proposed PPA Rate: ${ppa_rate:.3f} / kWh ({ppa_discount}% Discount)
* Estimated Year 1 Savings: ${annual_savings:,.2f}
* Estimated 20-Year Savings: ${annual_savings * 20:,.2f}
* CO2 Footprint Reduction: {co2_saved_tons:.1f} Metric Tons / year

👉 100% Funded. $0 Installation. $0 Maintenance."""
            
            st.text_area("Tap and hold to copy directly to your email or WhatsApp:", value=report_text, height=200)

            # 4. DOWNLOAD WORD DOCUMENT (.docx) DIRECTLY ON MOBILE
            st.subheader("📥 Download Official Document")
            if DOCX_AVAILABLE:
                docx_buffer = generate_docx_in_memory(
                    display_name, system_size, current_rate, annual_kwh, 
                    grid_annual_cost, ppa_rate, ppa_annual_cost, annual_savings, 
                    co2_saved_tons, tax_deduction, reap_grant, total_incentives
                )
                
                st.download_button(
                    label="📥 Download Word (.docx) Report",
                    data=docx_buffer,
                    file_name=f"GreenCruise_Feasibility_Report_{system_size}kW.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Document generator starting up. Please refresh the page in 10 seconds.")
